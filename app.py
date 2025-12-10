import streamlit as st
import google.generativeai as genai
from docx import Document
import requests
from bs4 import BeautifulSoup
from io import BytesIO  # This helps us create files in memory

# --- CONFIGURATION ---
# 1. PASTE YOUR API KEY HERE
api_key = st.secrets["GOOGLE_API_KEY"]

# 2. Setup Google AI
# We use 'gemini-1.5-flash' as it is the standard. 
# If this fails, try changing it to 'gemini-pro'
try:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.0-flash')
except Exception as e:
    st.error(f"Error setting up AI. Details: {e}")

# --- HELPER FUNCTIONS ---

def read_word_doc(file):
    """Reads text from the uploaded Word file"""
    try:
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except:
        return "Error reading document."

def create_docx(text):
    """Converts the AI text back into a Word Document for download"""
    doc = Document()
    # Add a title
    doc.add_heading('Optimized Resume', 0)
    
    # Add the AI generated text
    # We split by new lines to keep paragraphs clean
    for line in text.split('\n'):
        doc.add_paragraph(line)
        
    # Save to a memory buffer (not hard drive yet)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def scrape_url(url):
    """Gets text from a job website"""
    try:
        response = requests.get(url, headers={"User-Agent": "Mozilla/5.0"})
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, "html.parser")
            text = soup.get_text(separator=' ')
            return " ".join(text.split())
        else:
            return None
    except:
        return None

# --- THE APP DESIGN ---

st.set_page_config(page_title="Resume Optimizer", page_icon="📄")
# --- CUSTOM DESIGN (CSS) ---
st.markdown("""
<style>
    /* Import a modern font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600&display=swap');

    /* Apply font to whole app */
    html, body, [class*="css"]  {
        font-family: 'Inter', sans-serif;
    }

    /* Style the main title */
    h1 {
        color: #4F8BF9; /* Electric Blue */
        font-weight: 700;
        text-align: center;
    }

    /* Style the buttons */
    .stButton>button {
        width: 100%;
        background-color: #4F8BF9;
        color: white;
        border-radius: 12px;
        height: 50px;
        font-size: 18px;
        font-weight: 600;
        border: none;
        box-shadow: 0px 4px 15px rgba(79, 139, 249, 0.4);
        transition: all 0.3s ease;
    }

    /* Button hover effect */
    .stButton>button:hover {
        background-color: #3b6ccf;
        transform: translateY(-2px);
    }

    /* Style the upload box */
    [data-testid='stFileUploader'] {
        background-color: #f0f2f6;
        border-radius: 15px;
        padding: 20px;
    }
    
    /* Center the download button */
    .stDownloadButton>button {
        width: 100%;
        background-color: #00CC66; /* Success Green */
        color: white;
        border-radius: 12px;
    }
</style>
""", unsafe_allow_html=True)
st.title("📄 AI Resume Optimizer")
st.markdown("Upload your current resume and a job description. The AI will rewrite it for you.")

# Sidebar (Inputs)
with st.sidebar:
    st.header("Step 1: Your Resume")
    uploaded_file = st.file_uploader("Upload Word Doc (.docx)", type=['docx'])
    
    st.header("Step 2: Job Description")
    jd_option = st.radio("Choose Input Method:", ["Paste Text", "Job Website URL"])
    
    jd_text = ""
    if jd_option == "Paste Text":
        jd_text = st.text_area("Paste Job Description Here", height=200)
    else:
        jd_url = st.text_input("Paste Job URL Here")
        if st.button("Fetch Job Data"):
            if jd_url:
                with st.spinner("Reading website..."):
                    fetched_data = scrape_url(jd_url)
                    if fetched_data:
                        jd_text = fetched_data
                        st.success("Job data loaded!")
                    else:
                        st.error("Could not read URL. Please paste text manually.")

    st.header("Step 3: Choose Mode")
    mode = st.radio("Goal:", ["Full Resume Rewrite", "ATS Optimization Check"])

# Main Logic
if st.button("✨ Optimize My CV"):
    # Error Checking
    if not uploaded_file:
        st.warning("Please upload a Resume.")
    elif not jd_text:
        st.warning("Please provide a Job Description.")
    # The Trap Check (To ensure you pasted the key)
    elif api_key == "PASTE_YOUR_GOOGLE_API_KEY_HERE":
        st.error("You forgot to paste your API Key in the code file!")
    else:
        # Read the resume
        resume_text = read_word_doc(uploaded_file)
        
        # Prepare the prompt
        final_prompt = ""
        if mode == "Full Resume Rewrite":
            final_prompt = f"""
            Act as an Expert Resume Writer.
            Generate a full resume draft tailored to this job description. 
            Use my existing resume as a base for experience and skills.
            Focus on highlighting the most relevant qualifications and incorporating keywords.
            
            My Resume: {resume_text}
            Job Description: {jd_text}
            """
        else:
            final_prompt = f"""
            Act as an ATS (Applicant Tracking System) Specialist.
            Optimize my current resume by incorporating relevant keywords from the job description.
            
            My Resume: {resume_text}
            Job Description: {jd_text}
            
            Provide output as:
            1. List of Missing Keywords
            2. Rewritten Bullet Points
            """

        # Ask AI
        with st.spinner("AI is writing your resume..."):
            try:
                response = model.generate_content(final_prompt)
                
                # Show Result on Screen
                st.subheader("Your Optimized Result")
                st.markdown(response.text)
                
                # --- NEW: DOWNLOAD BUTTON ---
                st.success("Optimization Complete! Download your file below.")
                
                # Convert the text to a Word Doc
                docx_file = create_docx(response.text)
                
                # Create the Download Button
                st.download_button(
                    label="📥 Download as Word Doc",
                    data=docx_file,
                    file_name="Optimized_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:

                st.error(f"AI Error: {e}")


